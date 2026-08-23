from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "tmp" / "cv" / "Azeez_Adewale_Hamzat_CV.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

INK = RGBColor(41, 44, 41)
SOFT = RGBColor(86, 91, 86)
GOLD = RGBColor(115, 89, 31)
LINE = "D8DAD3"
FONT = "Arial"


def set_font(run, size=None, color=INK, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def add_hyperlink(paragraph, text, url, color=GOLD):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), FONT)
    run_fonts.set(qn("w:hAnsi"), FONT)
    properties.append(run_fonts)
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), str(color))
    properties.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    properties.append(size)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rule(paragraph, color=LINE, size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = set_font(paragraph.add_run("Azeez Adewale Hamzat | CV | "), 8.5, SOFT)
    prefix.bold = False
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_section_heading(document, text):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.add_run(text)
    return paragraph


def add_entry(document, title, dates, organisation=None, note=None, bullets=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_after = Pt(1)
    set_font(paragraph.add_run(title), 10, INK, bold=True)
    date_run = set_font(paragraph.add_run(f"  |  {dates}"), 9.5, GOLD, bold=True)
    date_run.italic = False
    if organisation:
        organisation_paragraph = document.add_paragraph()
        organisation_paragraph.paragraph_format.keep_with_next = bool(note or bullets)
        organisation_paragraph.paragraph_format.space_after = Pt(2)
        set_font(organisation_paragraph.add_run(organisation), 9.5, SOFT, italic=True)
    if note:
        note_paragraph = document.add_paragraph(note)
        note_paragraph.paragraph_format.space_after = Pt(3)
    for bullet in bullets or []:
        bullet_paragraph = document.add_paragraph(style="List Bullet")
        bullet_paragraph.paragraph_format.keep_together = True
        bullet_paragraph.add_run(bullet)


document = Document()
properties = document.core_properties
properties.title = "Curriculum Vitae | Azeez Adewale Hamzat"
properties.subject = "Academic curriculum vitae"
properties.author = "Azeez Adewale Hamzat"
properties.last_modified_by = ""
properties.keywords = "collective intelligence, group decision-making, agricultural knowledge systems"
properties.comments = "Public academic curriculum vitae"
properties.language = "en-GB"

section = document.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.68)
section.right_margin = Inches(0.68)
section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.28)

styles = document.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(9.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.12

title_style = styles["Title"]
title_style.font.name = "Georgia"
title_style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
title_style.font.size = Pt(24)
title_style.font.bold = False
title_style.font.color.rgb = INK
title_style.paragraph_format.space_before = Pt(0)
title_style.paragraph_format.space_after = Pt(2)

heading = styles["Heading 1"]
heading.font.name = FONT
heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
heading.font.size = Pt(10.5)
heading.font.bold = True
heading.font.color.rgb = GOLD
heading.font.all_caps = True
heading.paragraph_format.space_before = Pt(8)
heading.paragraph_format.space_after = Pt(4)
heading.paragraph_format.keep_with_next = True

list_style = styles["List Bullet"]
list_style.font.name = FONT
list_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
list_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
list_style.font.size = Pt(9.5)
list_style.paragraph_format.left_indent = Inches(0.34)
list_style.paragraph_format.first_line_indent = Inches(-0.17)
list_style.paragraph_format.space_after = Pt(1.5)
list_style.paragraph_format.line_spacing = 1.08

footer = section.footer.paragraphs[0]
add_page_number(footer)

name = document.add_paragraph(style="Title")
name.add_run("Azeez Adewale Hamzat")

role = document.add_paragraph()
role.paragraph_format.space_after = Pt(5)
set_font(role.add_run("Doctoral Researcher in Collective Intelligence and Group Decision-Making"), 10.5, GOLD, bold=True)

contact = document.add_paragraph()
contact.paragraph_format.space_after = Pt(8)
set_font(contact.add_run("Dublin, Ireland  |  "), 9.5, SOFT)
add_hyperlink(contact, "Email", "mailto:azeezhamzat@yahoo.com")
set_font(contact.add_run("  |  "), 9.5, SOFT)
add_hyperlink(contact, "Website", "https://azeezhamzat.com")
set_font(contact.add_run("  |  "), 9.5, SOFT)
add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/azeezhamzat")
set_font(contact.add_run("  |  "), 9.5, SOFT)
add_hyperlink(contact, "ORCID", "https://orcid.org/0009-0007-8804-0283")
set_font(contact.add_run("  |  "), 9.5, SOFT)
add_hyperlink(contact, "Google Scholar", "https://scholar.google.com/citations?user=6CErkzgAAAAJ")
add_rule(contact)

add_section_heading(document, "Research interests")
document.add_paragraph(
    "Collective intelligence and group decision-making; information diversity and collective accuracy; agricultural knowledge systems; participatory decision-making; human-AI interaction in groups; agent-based modelling of social dynamics."
)

add_section_heading(document, "Education")
add_entry(
    document,
    "PhD Researcher, Collective Intelligence and Group Decision-Making",
    "2025-present",
    "TCD-TUD SOHAM Centre, Technological University Dublin, Ireland",
    "Doctoral fellowship investigating information diversity, influence, and collective accuracy when AI becomes part of group decision processes.",
)
add_entry(
    document,
    "MSc, Collective Intelligence",
    "2020-2022",
    "Mohammed VI Polytechnic University (UM6P), Morocco",
    "Ibn Rochd Excellence Scholarship. Thesis on smallholder farmers' perceptions of cluster farming in Nigeria using a community knowledge-based assessment.",
)
add_entry(
    document,
    "BTech, Crop and Environmental Protection",
    "2009-2014",
    "Ladoke Akintola University of Technology (LAUTECH), Nigeria",
    "Foundation in agricultural science, environmental systems, and field research.",
)

add_section_heading(document, "Publications and research outputs")
add_entry(
    document,
    "Four Decades of Greenhouse Gas Emissions in Africa (1970-2024): Structural Trends, Sub-Regional Divergence, and the Case for Climate Equity",
    "Submitted manuscript, 2026",
    "Hamzat, A. A. | Independent research",
    "Details withheld while the editorial process is ongoing.",
)
add_entry(
    document,
    "Towards Anticipatory Innovation Systems: Learning from Africa's Triangular Trap",
    "Forthcoming",
    "Karuri-Sebina, G., Adesida, O., and Hamzat, A. A. | Accepted book chapter",
    "In The New Handbook for National Innovation Systems and Developing Countries. Publication details forthcoming.",
)
publication = document.add_paragraph()
publication.paragraph_format.space_after = Pt(4)
set_font(publication.add_run("Hamzat, A. A. (2022). "), 9.5, INK)
set_font(publication.add_run("A Community Knowledge-Based Assessment of Smallholder Farmers' Perception of Cluster Farming. "), 9.5, INK, italic=True)
add_hyperlink(publication, "Zenodo preprint", "https://doi.org/10.5281/zenodo.7678799")

add_section_heading(document, "Professional experience")
add_entry(
    document,
    "Outreach and Communications Officer",
    "2023-2025",
    "Africa Initiative, Mohammed VI Polytechnic University (UM6P), Morocco",
    bullets=[
        "Designed outreach strategies connecting UM6P with educational institutions across Africa.",
        "Facilitated collaborative sessions, co-organised events, and supported cross-institutional meetings.",
        "Created and maintained the department website and produced more than 100 weekly newsletter editions.",
        "Applied collective-intelligence principles to organisational communication and stakeholder engagement.",
    ],
)
add_entry(
    document,
    "Packaging Supervisor",
    "2017-2020",
    "Olam International, Nigeria",
    bullets=[
        "Supervised production-line staff and equipment across multiple shifts.",
        "Coordinated quality-control processes and resolved production bottlenecks under tight deadlines.",
    ],
)

document.add_page_break()

add_section_heading(document, "Research experience")
add_entry(
    document,
    "Master's Thesis Research",
    "2021-2022",
    "UM6P and fieldwork in Nigeria",
    bullets=[
        "Designed and conducted a community knowledge-based assessment of smallholder farmer perceptions.",
        "Collected and analysed qualitative and quantitative evidence from farming communities in southwestern Nigeria.",
        "Connected agricultural extension methods with collective-intelligence frameworks.",
    ],
)
add_entry(
    document,
    "Computational Research Projects",
    "2020-2022",
    "Mohammed VI Polytechnic University, Morocco",
    bullets=[
        "Built an agent-based model of misinformation evolution in social networks.",
        "Conducted an exploratory corpus-based analysis of 607 climate-technology articles using word embeddings, sentiment analysis, and factor analysis, with comparison to global emissions data.",
    ],
)

add_section_heading(document, "Independent project")
add_entry(
    document,
    "Constellation, Founder and Developer",
    "2026-present",
    "Private collective-intelligence platform",
    "Designed and developed a private environment for trusted groups to coordinate support, record decisions, and preserve shared learning without turning relationships into public performance metrics.",
)

add_section_heading(document, "Methods and technical skills")
add_entry(document, "Programming and data", "", note="Python, R, agent-based modelling, data analysis, data visualisation, and natural language processing.")
add_entry(document, "Research methods", "", note="Participatory research, fieldwork, survey design, qualitative analysis, community knowledge assessment, and experimental thinking.")
add_entry(document, "Communication", "", note="Academic writing, science communication, stakeholder engagement, facilitation, and web development.")

add_section_heading(document, "Languages")
document.add_paragraph("English and Yoruba (native); French (working proficiency); Arabic (basic).")

add_section_heading(document, "Awards and fellowships")
add_entry(document, "Doctoral Fellowship", "2025", "TCD-TUD SOHAM Centre")
add_entry(document, "Ibn Rochd Excellence Scholarship", "2020", "Mohammed VI Polytechnic University")

document.save(OUT)
print(OUT)
